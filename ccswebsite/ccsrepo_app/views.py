from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth import password_validation
from django.core.exceptions import ValidationError
from django.utils.translation import gettext as _
from django.utils.timezone import now
from django.contrib import messages
from django.utils import timezone
import os
from django.db.models import Q 
import pytesseract
from .models import CustomUser, Program, Category, ManuscriptType, Batch, AdviserStudentRelationship, Manuscript, PageOCRData, ManuscriptAccessRequest, Keyword, ManuscriptView
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from io import BytesIO
from django.db.models import Count
from django.http import JsonResponse
from django.utils.html import mark_safe
import re
from PIL import Image
import fitz
from django.core.exceptions import ObjectDoesNotExist
from django.db import IntegrityError
from django.db.models import Case, When, Value, IntegerField
from docx import Document
from django.http import HttpResponse
from django.db.models import Sum
#----------------Search and Manuscript flow System ------------------------/
def get_filtered_manuscripts(search_query, program_id=None, manuscript_type_id=None, category_id=None):
    # Get all approved manuscripts
    manuscripts = Manuscript.objects.filter(status='approved')

    # Filter based on search query for title, abstracts, batch name, and category
    if search_query:
        manuscripts = manuscripts.filter(
            Q(title__icontains=search_query) |
            Q(abstracts__icontains=search_query) |
            Q(year__icontains=search_query) |
            Q(category__name__icontains=search_query) |
            Q(adviser__first_name__icontains=search_query) |
            Q(adviser__last_name__icontains=search_query) |
            Q(authors__icontains=search_query) |
            Q(program__name__icontains=search_query) |
            Q(manuscript_type__name__icontains=search_query)
        )

    # Filter by program, manuscript type, category, and batch if specified
    if program_id:
        manuscripts = manuscripts.filter(program_id=program_id)
    if manuscript_type_id:
        manuscripts = manuscripts.filter(manuscript_type_id=manuscript_type_id)
    if category_id:
        manuscripts = manuscripts.filter(category_id=category_id)

    return manuscripts.order_by('-publication_date')

def manuscript_search_page(request):
    search_query = request.GET.get('search', '')
    program_id = request.GET.get('program')
    manuscript_type_id = request.GET.get('manuscript_type')
    category_id = request.GET.get('category')

    # Get filtered manuscripts based on search query and filters
    manuscripts = get_filtered_manuscripts(search_query, program_id, manuscript_type_id, category_id)

    # Retrieve additional filter options
    programs = Program.objects.all()
    manuscript_types = ManuscriptType.objects.all()
    categories = Category.objects.all()

    # Pagination
    paginator = Paginator(manuscripts, 5)
    page_number = request.GET.get('page')
    manuscripts = paginator.get_page(page_number)

    return render(request, 'ccsrepo_app/manuscript_search_page.html', {
        'manuscripts': manuscripts,
        'search_query': search_query,
        'programs': programs,
        'manuscript_types': manuscript_types,
        'categories': categories,
    })

def view_manuscript(request, manuscript_id):
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)

    # Check if the user is the student or adviser of the manuscript
    is_student = request.user == manuscript.student
    
    # Check if the user is an admin
    is_admin = request.user.is_admin
    is_adviser = request.user.is_adviser

    # Check if the user has an approved access request
    access_request = ManuscriptAccessRequest.objects.filter(
        manuscript=manuscript,
        student=request.user,
        status='approved',
        access_start_date__lte=timezone.now(),
        access_end_date__gte=timezone.now()
    ).first()
    
    # Check if the user has a pending access request
    has_pending_request = ManuscriptAccessRequest.objects.filter(
        manuscript=manuscript,
        student=request.user,
        status='pending'
    ).exists()

    # Set has_access to True if the user is the student, adviser, admin, or has an approved request
    has_access = is_student or is_adviser or is_admin or (access_request is not None)

    return render(request, 'ccsrepo_app/view_manuscript.html', {
        'manuscript': manuscript,
        'has_access': has_access,
        'has_pending_request': has_pending_request,
    })

#----------------End Search ------------------------/
#dashboard
@login_required(login_url='login')
def dashboard_view(request):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)
    return render(request, 'ccsrepo_app/dashboard.html')

#logout
def logout_view(request):
    logout(request)
    return redirect('home')

#Logging In
def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        # Check if a user with the given username exists
        try:
            user = CustomUser.objects.get(username=username)
        except CustomUser.DoesNotExist:
            messages.error(request, "Invalid username or password. Please try again.")
            return redirect('login')

        # Check if the user is active
        if not user.is_active:
            messages.error(request, "Your account is inactive. Please verify your email to activate your account.")
            return redirect('login')

        # Authenticate the user
        user = authenticate(request, username=username, password=password)

        if user is not None:
            # Log in the user
            login(request, user)

            # Redirect based on user type
            if user.is_student:
                return redirect('visitor_search_manuscripts')
            elif user.is_adviser:
                return redirect('adviser_approve_student')
            elif user.is_admin:
                return redirect('manage_users')

            return redirect('adviser_request')
        else:
            messages.error(request, "Invalid username or password. Please try again.")
            return redirect('login')

    return render(request, 'ccsrepo_app/login.html')

#----------------Student and Adviser ------------------------/
#View of Request to Adviser
def success_request_view(request):
    return render(request, 'ccsrepo_app/adviser_request_success.html')

@login_required(login_url='login')
def request_adviser_view(request):
    print(f"Current user: {request.user}")

    if request.method == 'POST':
        adviser_email = request.POST.get('email')
        student = request.user 

        try:
            adviser = CustomUser.objects.get(email=adviser_email, is_adviser=True)

            # Check if there's already an approved relationship with the adviser
            existing_relationship = AdviserStudentRelationship.objects.filter(adviser=adviser, student=student).first()

            if existing_relationship:
                if existing_relationship.status == 'approved':
                    messages.warning(request, "You have already been approved by this adviser.")
                else:
                    messages.warning(request, "You have already sent a request to this adviser.")
            else:
                # Create a new adviser-student relationship with 'pending' status
                AdviserStudentRelationship.objects.create(adviser=adviser, student=student, status='pending')
                return redirect('adviser_request_success')

        except CustomUser.DoesNotExist:
            messages.error(request, "No adviser found with this email or they are not an adviser.")
        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")

    return render(request, 'ccsrepo_app/adviser_request.html')

@login_required(login_url='login')
def approve_student_view(request):
    if not request.user.is_adviser:
        return render(request, 'unauthorized.html', status=403)

    # Get all relationships where the logged-in user is the adviser, ordered by created_at
    relationships = AdviserStudentRelationship.objects.filter(adviser=request.user).order_by('-created_at')

    # Pagination setup: Show 5 relationships per page
    paginator = Paginator(relationships, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    if request.method == 'POST':
        student_id = request.POST.get('student_id')
        try:
            # Get the student relationship for the adviser
            student_relationship = get_object_or_404(AdviserStudentRelationship, id=student_id, adviser=request.user)
            
            # Get the student from the relationship
            student = student_relationship.student

            # Check if the status is already approved before updating
            if student_relationship.status != 'approved':
                # Update the status of the relationship to approved
                student_relationship.status = AdviserStudentRelationship.APPROVED  # Assuming you added the constant 'APPROVED'
                student_relationship.save()

                # Optionally, set the student role to True if needed (e.g., if your CustomUser model has is_student)
                student.is_student = True
                student.save()
            else:
                messages.info(request, f"{student.username} has already been approved.")

            return redirect('adviser_approve_student')

        except Exception as e:
            messages.error(request, f"An error occurred: {str(e)}")

    return render(request, 'ccsrepo_app/adviser_approve_student.html', {'page_obj': page_obj})

#----------------End Student and Adviser ------------------------/
#----------------Register ------------------------/
def validate_user_password(password):
    errors = []
    if len(password) < 8:
        errors.append(_("Password must be at least 8 characters long."))
    if not any(char.isupper() for char in password):
        errors.append(_("Password must contain at least one uppercase letter."))
    if not any(char.islower() for char in password):
        errors.append(_("Password must contain at least one lowercase letter."))
    if not any(char.isdigit() for char in password):
        errors.append(_("Password must contain at least one digit."))
    if not any(char in '!@#$%^&*()-_=+[]{}|;:,.<>?/' for char in password):
        errors.append(_("Password must contain at least one special character."))
    
    return errors

def validate_user_data(email, username, password1, password2):
    errors = {}

    # Updated email pattern: two letters, a four-digit year, five digits, and the domain
    email_pattern = r"^[a-zA-Z]{2}\d{4}\d{5}@wmsu\.edu\.ph$"
    if not re.match(email_pattern, email):
        errors['email'] = [_("Email must be wmsu email")]

    # Validate password with custom rules
    password_errors = validate_user_password(password1)
    if password_errors:
        errors['password1'] = password_errors

    # Check if passwords match
    if password1 != password2:
        errors['password2'] = [_("Passwords do not match.")]

    # Validate using Django's built-in password validators
    try:
        password_validation.validate_password(password1)
    except ValidationError as e:
        errors['password1'] = errors.get('password1', []) + list(e.messages)

    # Check if email or username already exists
    if CustomUser.objects.filter(email=email).exists():
        errors['email'] = errors.get('email', []) + [_("Email already exists.")]
    if CustomUser.objects.filter(username=username).exists():
        errors['username'] = [_("Username already exists.")]

    return errors

from django.contrib.sites.shortcuts import get_current_site
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from django.contrib.auth.tokens import default_token_generator
from django.template.loader import render_to_string
from django.core.mail import send_mail
from django.conf import settings

def StudentRegister(request):
    if request.method == 'POST':
        # Retrieve and strip form data
        email = request.POST.get('email', '').strip()
        username = request.POST.get('username', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        middle_name = request.POST.get('middle_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        password1 = request.POST.get('password1', '').strip()
        password2 = request.POST.get('password2', '').strip()
        program_id = request.POST.get('program')

        # Validate user data
        errors = validate_user_data(email, username, password1, password2)

        if errors:
            programs = Program.objects.all()
            return render(request, 'ccsrepo_app/register.html', {
                'programs': programs,
                'errors': errors,
                'email': email,
                'username': username,
                'first_name': first_name,
                'middle_name': middle_name,
                'last_name': last_name,
                'program_id': program_id
            })

        # Create the user but set `is_active` to False until email verification
        user = CustomUser(
            email=email,
            username=username,
            first_name=first_name,
            middle_name=middle_name,
            last_name=last_name,
            is_student=False,
            is_active=False,  # Account is inactive until verified
            program_id=program_id
        )
        user.set_password(password1)
        user.save()

        # Send activation email
        current_site = get_current_site(request)
        subject = 'Activate Your Account'
        message = render_to_string('ccsrepo_app/activation_email.html', {
            'user': user,
            'domain': current_site.domain,
            'uid': urlsafe_base64_encode(force_bytes(user.pk)),
            'token': default_token_generator.make_token(user),
        })
        send_mail(subject, message, settings.EMAIL_HOST_USER, [user.email], fail_silently=False)

        return render(request, 'ccsrepo_app/email_sent.html', {'email': email})
    
    programs = Program.objects.all()
    return render(request, 'ccsrepo_app/register.html', {'programs': programs})

def activate_account(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = CustomUser.objects.get(pk=uid)
    except (TypeError, ValueError, OverflowError, CustomUser.DoesNotExist):
        user = None

    if user and default_token_generator.check_token(user, token):
        user.is_active = True
        user.save()
        return redirect('login')
    else:
        return render(request, 'ccsrepo_app/activation_invalid.html')
#----------------End Register------------------------/

#----------------Admin Managing ------------------------/
#Program
@login_required(login_url='login')
def manage_program(request):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)
    if request.method == 'POST':
        name = request.POST.get('name')
        abbreviation = request.POST.get('abbreviation')
        
        if Program.objects.filter(name=name).exists():
            messages.error(request, "A program with this name already exists.")
            return redirect('manage_program')
        
        Program.objects.create(name=name, abbreviation=abbreviation)
        return redirect('manage_program')

    programs = Program.objects.all()
    paginator = Paginator(programs, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'ccsrepo_app/manage_program.html', {
        'programs': programs,
        'page_obj': page_obj
    })

@login_required(login_url='login')
def dashboard_page(request):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)

    # Manuscripts by category
    category_counts = (
        Manuscript.objects.values('category__name')
        .annotate(count=Count('id'))
        .order_by('-count')
    )
    category_data = {
        'labels': [item['category__name'] for item in category_counts if item['category__name']],
        'data': [item['count'] for item in category_counts if item['category__name']],
    }

    # Manuscripts by program
    program_counts = (
        Manuscript.objects.values('program__name')
        .annotate(count=Count('id'))
        .order_by('-count')
    )
    program_data = {
        'labels': [item['program__name'] for item in program_counts if item['program__name']],
        'data': [item['count'] for item in program_counts if item['program__name']],
    }

    # Manuscripts by type
    type_counts = (
        Manuscript.objects.values('manuscript_type__name')
        .annotate(count=Count('id'))
        .order_by('-count')
    )
    type_data = {
        'labels': [item['manuscript_type__name'] for item in type_counts if item['manuscript_type__name']],
        'data': [item['count'] for item in type_counts if item['manuscript_type__name']],
    }

    # Data for summary cards
    program_summary = [{'name': item['program__name'], 'count': item['count']} for item in program_counts if item['program__name']]
    category_summary = [{'name': item['category__name'], 'count': item['count']} for item in category_counts if item['category__name']]
    type_summary = [{'name': item['manuscript_type__name'], 'count': item['count']} for item in type_counts if item['manuscript_type__name']]

    return render(request, 'ccsrepo_app/dashboard_page.html', {
        'category_data': category_data,
        'program_data': program_data,
        'type_data': type_data,
        'program_summary': program_summary,
        'category_summary': category_summary,
        'type_summary': type_summary,
    })

@login_required(login_url='login')
def manage_category(request):
    if request.user.is_admin:
        if request.method == 'POST':
            name = request.POST.get('name')

            if Category.objects.filter(name=name).exists():
                messages.error(request, "A category with this name already exists.")
                return redirect('manage_category')

            Category.objects.create(name=name)
            messages.success(request, "Category added successfully.")
            return redirect('manage_category')

        category = Category.objects.all()
        paginator = Paginator(category, 5)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        return render(request, 'ccsrepo_app/manage_category.html', {
            'category': category,
            'page_obj': page_obj
        })
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def edit_category(request, category_id):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)

    category = get_object_or_404(Category, id=category_id)

    if request.method == 'POST':
        name = request.POST.get('name')

        if Category.objects.filter(name=name).exists() and name != category.name:
            messages.error(request, "A category with this name already exists.")
            return redirect('edit_category', category_id=category.id)

        category.name = name
        category.save()
        messages.success(request, "Category updated successfully.")
        return redirect('manage_category')

    return render(request, 'ccsrepo_app/edit_category.html', {'category': category})


@login_required(login_url='login')
def edit_program(request, program_id):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)

    program = get_object_or_404(Program, id=program_id)

    if request.method == 'POST':
        name = request.POST.get('name').strip()
        abbreviation = request.POST.get('abbreviation').strip()

        if Program.objects.filter(name__iexact=name).exclude(id=program_id).exists():
            messages.error(request, "A program with this name already exists.")
            return redirect('edit_program', program_id=program.id)

        if Program.objects.filter(abbreviation__iexact=abbreviation).exclude(id=program_id).exists():
            messages.error(request, "A program with this abbreviation already exists.")
            return redirect('edit_program', program_id=program.id)

        program.name = name
        program.abbreviation = abbreviation
        program.save()
        messages.success(request, "Program updated successfully.")
        return redirect('manage_program')

    return render(request, 'ccsrepo_app/edit_program.html', {'program': program})


@login_required(login_url='login')
def edit_type(request, type_id):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)

    manuscript_type = get_object_or_404(ManuscriptType, id=type_id)

    if request.method == 'POST':
        name = request.POST.get('name')

        if ManuscriptType.objects.filter(name=name).exists() and name != manuscript_type.name:
            messages.error(request, "A manuscript type with this name already exists.")
            return redirect('edit_type', type_id=manuscript_type.id)

        manuscript_type.name = name
        manuscript_type.save()
        messages.success(request, "Manuscript type updated successfully.")
        return redirect('manage_type')

    return render(request, 'ccsrepo_app/edit_type.html', {'type': manuscript_type})


@login_required(login_url='login')
def edit_adviser(request, adviser_id):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)

    adviser = get_object_or_404(CustomUser, id=adviser_id)
    programs = Program.objects.all()

    if request.method == 'POST':
        first_name = request.POST.get('first_name').strip()
        middle_name = request.POST.get('middle_name').strip()
        last_name = request.POST.get('last_name').strip()
        program_id = request.POST.get('program')
        username = request.POST.get('username').strip()
        email = request.POST.get('email').strip()
        password1 = request.POST.get('password1').strip()
        password2 = request.POST.get('password2').strip()

        if password1 and password1 != password2:
            messages.error(request, "Passwords do not match.")
            return render(request, 'ccsrepo_app/edit_adviser.html', {
                'adviser': adviser,
                'programs': programs,
                'errors': {'password1': ['Passwords do not match.']}
            })

        adviser.first_name = first_name
        adviser.middle_name = middle_name
        adviser.last_name = last_name
        adviser.username = username
        adviser.email = email

        if program_id:
            adviser.program = get_object_or_404(Program, id=program_id)

        if password1:
            adviser.set_password(password1)

        adviser.save()
        messages.success(request, "Adviser updated successfully.")
        return redirect('manage_users')

    return render(request, 'ccsrepo_app/edit_adviser.html', {
        'adviser': adviser,
        'programs': programs,
    })

@login_required(login_url='login')
def manage_batch(request):
    if request.user.is_admin:
        if request.method == 'POST':
            name = request.POST.get('name', '').strip()

            # Check if the batch name already exists
            if Batch.objects.filter(name__iexact=name).exists():
                messages.error(request, "A batch with this name already exists.")
            else:
                # Create the batch if it doesn't already exist
                Batch.objects.create(name=name)
                messages.success(request, "Batch created successfully.")
            return redirect('manage_batch')

        # Fetch all batches to display in the template
        batch = Batch.objects.all()
        return render(request, 'ccsrepo_app/manage_batch.html', {'batch': batch})

    # Render the unauthorized page for non-admin users
    return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def manage_type(request):
    if request.user.is_admin:
        if request.method == 'POST':
            name = request.POST.get('name', '').strip()

            # Check if the type name already exists
            if ManuscriptType.objects.filter(name__iexact=name).exists():
                messages.error(request, "A type with this name already exists.")
            else:
                # Create the type if it doesn't already exist
                ManuscriptType.objects.create(name=name)
                messages.success(request, "Type created successfully.")
            return redirect('manage_type')

        # Paginate manuscript types
        type_list = ManuscriptType.objects.all()
        paginator = Paginator(type_list, 5)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'ccsrepo_app/manage_type.html', {
            'type': type_list,
            'page_obj': page_obj
        })

    # Render the unauthorized page for non-admin users
    return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def ManageAdviser(request):
    if request.user.is_admin:
        if request.method == 'POST':
            email = request.POST.get('email', '').strip()
            username = request.POST.get('username', '').strip()
            first_name = request.POST.get('first_name', '').strip()
            middle_name = request.POST.get('middle_name', '').strip() or None  # Allow empty or None if not provided
            last_name = request.POST.get('last_name', '').strip()
            password1 = request.POST.get('password1', '').strip()
            password2 = request.POST.get('password2', '').strip()
            program_id = request.POST.get('program', '').strip()

            # Validate user data
            errors = validate_adviser_data(email, username, first_name, middle_name, last_name, program_id, password1, password2)

            if errors:
                # Reload programs and advisers to re-render the form with errors
                programs = Program.objects.all()
                advisers = CustomUser.objects.filter(is_adviser=True)
                return render(request, 'ccsrepo_app/manage_users.html', {
                    'programs': programs,
                    'advisers': advisers,
                    'errors': errors,
                    'email': email,
                    'username': username,
                    'first_name': first_name,
                    'middle_name': middle_name,
                    'last_name': last_name,
                    'program_id': program_id
                })

            # Create adviser if validation passes
            user = CustomUser(
                email=email,
                username=username,
                first_name=first_name,
                middle_name=middle_name,  # Middle name can be None or empty
                last_name=last_name,
                is_adviser=True,
                program_id=program_id
            )
            user.set_password(password1)
            user.save()
            messages.success(request, "Adviser created successfully.")
            return redirect('manage_adviser')

        # Fetch programs and advisers for rendering
        programs = Program.objects.all()
        advisers = CustomUser.objects.filter(is_adviser=True)

        # Paginate advisers
        paginator = Paginator(advisers, 5)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'ccsrepo_app/manage_users.html', {
            'advisers': advisers,
            'page_obj': page_obj,
            'programs': programs
        })

    # Render the unauthorized page for non-admin users
    return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def create_program(request):
    if request.user.is_admin:
        # Initialize variables to retain form values in case of validation failure
        name = ""
        abbreviation = ""

        if request.method == "POST":
            # Retrieve and strip form inputs
            name = request.POST.get("name", "").strip()
            abbreviation = request.POST.get("abbreviation", "").strip()

            # Check for duplicates in name and abbreviation
            if Program.objects.filter(name__iexact=name).exists():
                messages.error(request, _("A program with this name already exists."))
            elif Program.objects.filter(abbreviation__iexact=abbreviation).exists():
                messages.error(request, _("A program with this abbreviation already exists."))
            else:
                # Create the program if no duplicates are found
                Program.objects.create(name=name, abbreviation=abbreviation)
                messages.success(request, _("Program created successfully."))
                return redirect("manage_program")  # Redirect to the program management page

        # Render the form for GET requests or on validation failure
        return render(request, "ccsrepo_app/create_program.html", {"program": {"name": name, "abbreviation": abbreviation}})

    # Render the unauthorized page for non-admin users
    return render(request, "unauthorized.html", status=403)

@login_required(login_url='login')
def create_category(request):
    if request.user.is_admin:
        # Initialize variable to retain form values in case of validation failure
        name = ""

        if request.method == "POST":
            # Retrieve and strip form inputs
            name = request.POST.get("name", "").strip()

            # Check for duplicate category name
            if Category.objects.filter(name__iexact=name).exists():
                messages.error(request, _("A category with this name already exists."))
            else:
                # Create the category if no duplicates are found
                Category.objects.create(name=name)
                messages.success(request, _("Category created successfully."))
                return redirect("manage_category")  # Redirect to the category management page

        # Render the form for GET requests or on validation failure
        return render(request, "ccsrepo_app/create_category.html", {"category": {"name": name}})

    # Render the unauthorized page for non-admin users
    return render(request, "unauthorized.html", status=403)

@login_required(login_url='login')
def create_manuscripttype(request):
    if request.user.is_admin:
        # Initialize an empty dictionary for storing the form values in case of errors
        manuscripttype = {'name': ''}

        if request.method == 'POST':
            # Retrieve and strip the name field from the form submission
            name = request.POST.get('name', '').strip()

            # Check if a manuscript type with the same name already exists (case-insensitive)
            if ManuscriptType.objects.filter(name__iexact=name).exists():
                messages.error(request, _("A manuscript type with this name already exists."))
                return render(request, 'ccsrepo_app/create_manuscripttype.html', {'manuscripttype': {'name': name}})

            # If no duplicate exists, create the new manuscript type
            ManuscriptType.objects.create(name=name)
            messages.success(request, _("Manuscript type created successfully."))
            return redirect('manage_type')  # Redirect to the page where manuscript types are listed

        # Render the form on a GET request or after validation failure
        return render(request, 'ccsrepo_app/create_manuscripttype.html', {'manuscripttype': manuscripttype})

    # Render the unauthorized page for unauthorized users
    return render(request, 'unauthorized.html', status=403)

def check_duplicate_manuscripttype(request):
    name = request.GET.get('name', '').strip()
    duplicate_name = ManuscriptType.objects.filter(name__iexact=name).exists()
    return JsonResponse({'duplicate_name': duplicate_name})

def check_duplicate_category(request):
    name = request.GET.get('name', '').strip()
    duplicate_name = Category.objects.filter(name__iexact=name).exists()
    return JsonResponse({'duplicate_name': duplicate_name})

def check_program_duplicate(request):
    name = request.GET.get("name", "").strip()
    abbreviation = request.GET.get("abbreviation", "").strip()

    duplicate_name = Program.objects.filter(name__iexact=name).exists()
    duplicate_abbreviation = Program.objects.filter(abbreviation__iexact=abbreviation).exists()

    return JsonResponse({
        "duplicate_name": duplicate_name,
        "duplicate_abbreviation": duplicate_abbreviation,
    })

@login_required(login_url='login')
def create_adviser(request):
    if not request.user.is_admin:
        return render(request, 'unauthorized.html', status=403)
    if request.method == 'POST':
        # Retrieve form data and process
        email = request.POST.get('email').strip()
        username = request.POST.get('username').strip()
        first_name = request.POST.get('first_name').strip()
        middle_name = request.POST.get('middle_name').strip()
        last_name = request.POST.get('last_name').strip()
        password1 = request.POST.get('password1').strip()
        password2 = request.POST.get('password2').strip()
        program_id = request.POST.get('program')

        # Validate and create user
        errors = validate_adviser_data(email, username, first_name, middle_name, last_name, program_id, password1, password2)

        if errors:
            # Return with errors
            programs = Program.objects.all()
            return render(request, 'ccsrepo_app/create_adviser.html', {
                'errors': errors,
                'programs': programs,
                'email': email,
                'username': username,
                'first_name': first_name,
                'middle_name': middle_name,
                'last_name': last_name,
                'program_id': program_id
            })
        
        # If validation passes, create and save the adviser
        user = CustomUser(
            email=email,
            username=username,
            first_name=first_name,
            middle_name=middle_name,
            last_name=last_name,
            is_adviser=True,
            program_id=program_id
        )
        user.set_password(password1)
        user.save()

        return redirect('manage_users')  # Redirect to the manage advisers page

    else:
        programs = Program.objects.all()
        return render(request, 'ccsrepo_app/create_adviser.html', {'programs': programs})
    
def validate_adviser_data(email, username, first_name, middle_name, last_name, program_id, password1, password2):
    errors = {}

    # Validate first name and last name
    if not first_name:
        errors['first_name'] = [_("First name is required.")]
    if not last_name:
        errors['last_name'] = [_("Last name is required.")]

    # Middle name is optional, no need to validate its presence
    if middle_name and len(middle_name) < 2:  # Optional validation, if provided, check its length or other conditions
        errors['middle_name'] = [_("Middle name is too short.")]

    # Validate program
    if not program_id:
        errors['program'] = [_("Program is required.")]

    # Password validation
    try:
        password_validation.validate_password(password1)
    except ValidationError as e:
        errors['password1'] = list(e.messages)

    # Updated email pattern: two letters, a four-digit year, five digits, and the domain
    email_pattern = r"^[a-zA-Z]{2}\d{4}\d{5}@wmsu\.edu\.ph$"
    if not re.match(email_pattern, email):
        errors['email'] = [_("Email must be a valid WMSU email address.")]

    # Check if passwords match
    if password1 != password2:
        errors['password2'] = [_("Passwords do not match.")]

    # Check if email already exists
    if CustomUser.objects.filter(email=email).exists():
        errors['email'] = [_("Email already exists.")]

    # Check if username already exists
    if CustomUser.objects.filter(username=username).exists():
        errors['username'] = [_("Username already exists.")]
    
    return errors
#----------------End Admin Managing ------------------------/


#----------------Manuscript System ------------------------/
#----------------UTIL/Helper------------------------/

def validate_user_title(title):
    errors = []
    
    # Check if any other manuscript with the same title has upload_show=True
    if Manuscript.objects.filter(title=title, upload_show=True).exists():
        errors.append(_("A manuscript with this title is already published (upload_show=True). Please choose a different title."))
    
    return errors

def extract_title_from_first_page(first_page_text):
    """Extract and format the title from the first page."""
    title_text = first_page_text.replace('\n', ' ').strip()
    title_text = re.sub(r'\s+', ' ', title_text)  # Normalize spaces to a single space

    # Define possible start and end keywords for the title extraction
    title_start_keywords = ["Zamboanga City", "Department of Information Technology"]
    title_end_keywords = [
        "A thesis presented to the faculty", "In partial fulfillment",
        "A capstone project", "A CAPSTONE PROJECT"
    ]
    title_exclude_keywords = [
        "A CAPSTONE PROJECT Presented to the Faculty of the College of Computing Studies Western Mindanao State University"
    ]

    start_idx = end_idx = -1
    for start_keyword in title_start_keywords:
        start_idx = title_text.lower().find(start_keyword.lower())
        if start_idx != -1:
            break
    for end_keyword in title_end_keywords:
        temp_end_idx = title_text.lower().find(end_keyword.lower())
        if temp_end_idx != -1:
            end_idx = temp_end_idx
            break
    for exclude_keyword in title_exclude_keywords:
        exclude_idx = title_text.lower().find(exclude_keyword.lower())
        if exclude_idx != -1:
            end_idx = exclude_idx  # Adjust end index to exclude unwanted part
            break

    if start_idx != -1 and end_idx != -1:
        title = title_text[start_idx + len(title_start_keywords[0]): end_idx].strip()
        
        unwanted_prefixes = ["INFORMATION TECHNOLOGY", "Information Technology"]
        for prefix in unwanted_prefixes:
            if title.lower().startswith(prefix.lower()):
                title = title[len(prefix):].strip()
                break
        title = title.replace("&", "").strip()
        return title if title else "No title found"
    else:
        return "No title found"

def extract_year_from_first_page(first_page_text):
    """Extract the year from the first page."""
    year_match = re.search(r'\b(19|20)\d{2}\b', first_page_text)
    return year_match.group(0) if year_match else "Year not found"

def extract_authors_from_first_page(first_page_text):
    """Extract authors from the first page."""
    author_end_keywords = ["Researchers", "Researcher"]
    author_start_keywords = [
        "Bachelor of Science in Computer Science",
        "presented to the faculty of department of computer science college of computing studies"
    ]
    
    author_end_idx = -1
    for end_keyword in author_end_keywords:
        temp_end_idx = first_page_text.lower().find(end_keyword.lower())
        if temp_end_idx != -1:
            author_end_idx = temp_end_idx
            break

    if author_end_idx != -1:
        author_start_idx = -1
        for start_keyword in author_start_keywords:
            temp_start_idx = first_page_text.lower().rfind(start_keyword.lower(), 0, author_end_idx)
            if temp_start_idx != -1:
                author_start_idx = temp_start_idx + len(start_keyword)
                break

        if author_start_idx != -1:
            authors_text = first_page_text[author_start_idx:author_end_idx].strip()
            author_pattern = re.compile(r'([A-Z][a-z]*\.? ?){2,}')
            authors = [match.group(0).strip() for match in author_pattern.finditer(authors_text)]

            disallowed_keywords = ["Science", "Studies", "Computer"]
            authors = [
                author for author in authors
                if not any(keyword in author for keyword in disallowed_keywords)
            ]
            if authors:
                return ', '.join(authors)

    by_index = first_page_text.lower().find("by")
    if by_index != -1:
        following_text = first_page_text[by_index + len("by"):].strip()
        lines = following_text.splitlines()
        raw_authors = [line.strip() for line in lines if line.strip()][:3]

        formatted_authors = []
        for raw_author in raw_authors:
            parts = raw_author.split(', ')
            if len(parts) == 2:
                surname, given_names = parts[0], parts[1]
                formatted_authors.append(f"{given_names} {surname}")
            else:
                formatted_authors.append(raw_author)
        if formatted_authors:
            return ', '.join(formatted_authors)

    return "No authors found"

def process_and_extract_manuscript_data(pdf_path, manuscript, max_abstract_pages=5, chunk_size=5, max_pages=5):
    """Process the manuscript by extracting abstract, OCR data, title, year, and authors from the first `max_pages`."""
    doc = fitz.open(pdf_path)
    ocr_data_list = []
    abstract_text = None
    first_page_text = None

    # Ensure we only process up to `max_pages`
    pages_to_process = min(max_pages, len(doc))

    # Process up to `max_pages`
    for page_num in range(pages_to_process):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=120)
        img = Image.open(BytesIO(pix.tobytes("png")))
        page_text = pytesseract.image_to_string(img).strip()

        # Capture the first page's text for title, year, and authors extraction
        if page_num == 0:
            first_page_text = page_text

        # Attempt to extract abstract from the first `max_abstract_pages`
        if abstract_text is None and page_num < max_abstract_pages:
            lower_text = page_text.lower()
            if lower_text.startswith("abstract") or lower_text.startswith("executive summary"):
                keyword = "abstract" if lower_text.startswith("abstract") else "executive summary"
                extracted_text = page_text[len(keyword):].strip()
                if "keywords" in extracted_text:
                    extracted_text = extracted_text.split("keywords")[0].strip()
                abstract_text = extracted_text or "No abstract found"

        # Add OCR data for this page
        ocr_data_list.append(PageOCRData(manuscript=manuscript, page_num=page_num + 1, text=page_text))

        # Save OCR data in chunks to reduce memory usage
        if len(ocr_data_list) >= chunk_size:
            PageOCRData.objects.bulk_create(ocr_data_list)
            ocr_data_list = []

    # Save any remaining OCR data
    if ocr_data_list:
        PageOCRData.objects.bulk_create(ocr_data_list)

    # Extract title, year, and authors from the first page
    title = extract_title_from_first_page(first_page_text)
    year = extract_year_from_first_page(first_page_text)
    authors = extract_authors_from_first_page(first_page_text)

    # Update manuscript details
    manuscript.title = title
    manuscript.year = year
    manuscript.authors = authors
    manuscript.abstracts = abstract_text or "No abstract found"
    manuscript.page_count = len(doc)  # Total number of pages in the document
    manuscript.current_page_count = pages_to_process  # Only processed up to `max_pages`
    manuscript.remaining_page = len(doc) - pages_to_process  # Remaining pages, if any
    manuscript.save()

@login_required(login_url='login')
def upload_manuscript(request):
    if request.user.is_student:
        if request.method == 'POST':
            pdf_file = request.FILES.get('pdf_file')

            if pdf_file:
                # Create and save the Manuscript instance
                manuscript = Manuscript(
                    pdf_file=pdf_file,
                    student=request.user,
                    abstracts="No abstract found"  # Default value for abstracts
                )
                manuscript.save()
                
                pdf_file_path = manuscript.pdf_file.path

                # Check if the PDF file exists and process it
                if os.path.exists(pdf_file_path):
                    try:
                        process_and_extract_manuscript_data(pdf_file_path, manuscript, max_abstract_pages=5, chunk_size=5, max_pages=5)
                    except Exception as e:
                        print(f"Error processing PDF: {e}")  # Log the error (can be enhanced)

                # Redirect to the final manuscript page
                return redirect('final_manuscript_page', manuscript_id=manuscript.id)

        # Render the manuscript upload page if the request is not POST
        return render(request, 'ccsrepo_app/manuscript_upload_page.html')

    # Render unauthorized page for non-students
    return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def final_manuscript_page(request, manuscript_id):
    if request.user.is_student:
        manuscript = get_object_or_404(Manuscript, id=manuscript_id, student=request.user)
        errors = []

        if request.method == 'POST':
            # Collect form data
            title = request.POST.get('title')
            abstracts = request.POST.get('abstracts')
            authors = request.POST.get('authors')
            year = request.POST.get('year')
            category_id = request.POST.get('category')
            manuscript_type_id = request.POST.get('manuscript_type')
            program_id = request.POST.get('program')
            adviser_id = request.POST.get('adviser')

            # Validate title uniqueness if upload_show=True
            if Manuscript.objects.filter(title=title, upload_show=True).exclude(id=manuscript.id).exists():
                errors.append(_("A manuscript with this title is already published. Please choose a different title."))

            # Validate year
            if not (year.isdigit() and len(year) == 4 and 1900 <= int(year) <= timezone.now().year):
                errors.append(_("Year must be a valid 4-digit number between 1900 and the current year."))

            # Validate adviser
            try:
                adviser = CustomUser.objects.get(id=adviser_id, is_adviser=True)
            except ObjectDoesNotExist:
                errors.append(_("Adviser not found. Please check the adviser ID."))

            # If there are errors, return them to the template
            if errors:
                categories = Category.objects.all()
                manuscript_types = ManuscriptType.objects.all()
                programs = Program.objects.all()
                advisers = CustomUser.objects.filter(is_adviser=True)
                return render(request, 'ccsrepo_app/manuscript_final_page.html', {
                    'manuscript': manuscript,
                    'categories': categories,
                    'manuscript_types': manuscript_types,
                    'programs': programs,
                    'advisers': advisers,
                    'errors': errors,
                })

            # Assign validated fields to the manuscript
            manuscript.title = title
            manuscript.abstracts = abstracts
            manuscript.authors = authors
            manuscript.year = year
            manuscript.category_id = category_id
            manuscript.manuscript_type_id = manuscript_type_id
            manuscript.program_id = program_id
            manuscript.adviser = adviser
            manuscript.upload_date = now()
            manuscript.upload_show = True
            manuscript.save()

            # **Keyword Extraction After Save**
            # Combine relevant fields for keyword extraction
            combined_text = f"{title} {abstracts} {authors}".lower()

            # Extract keywords based on tech-related terms
            extracted_keywords = extract_keywords_from_text(combined_text)

            # Save extracted keywords to the database
            existing_keywords = set(Keyword.objects.filter(manuscript=manuscript).values_list('keyword', flat=True))
            for keyword in extracted_keywords:
                if keyword not in existing_keywords:
                    Keyword.objects.create(manuscript=manuscript, keyword=keyword)

            return redirect('visitor_search_manuscripts')

        # Load choices for form in GET request
        categories = Category.objects.all()
        manuscript_types = ManuscriptType.objects.all()
        programs = Program.objects.all()
        advisers = CustomUser.objects.filter(is_adviser=True)

        return render(request, 'ccsrepo_app/manuscript_final_page.html', {
            'manuscript': manuscript,
            'categories': categories,
            'manuscript_types': manuscript_types,
            'programs': programs,
            'advisers': advisers,
        })

    # Render unauthorized page for non-students
    return render(request, 'unauthorized.html', status=403)


#----------------End Manuscript System ------------------------/

#----------------Adviser System ------------------------/
# def adviser_manuscript(request):
#     manuscripts = Manuscript.objects.filter(adviser=request.user).exclude(student=request.user)

#     return render(request, 'ccsrepo_app/adviser_manuscript.html', {
#         'manuscripts': manuscripts,
#     })
@login_required(login_url='login')
def adviser_manuscript(request):
    if request.user.is_adviser:
        # Get manuscripts for the adviser, including all statuses
        manuscripts = Manuscript.objects.filter(
            adviser=request.user
        ).exclude(
            student=request.user
        ).annotate(
            # Assign priority values for statuses
            status_priority=Case(
                When(status="pending", then=Value(1)),  # Pending manuscripts first
                When(status="review", then=Value(2)),  # Review manuscripts next
                When(status="approved", then=Value(3)),  # Approved manuscripts last
                default=Value(4),  # Fallback for any other statuses
                output_field=IntegerField(),
            )
        ).order_by(
            'status_priority',  # Sort by priority (pending first)
            '-upload_date'  # Within each status, sort by most recent upload date
        )

        # Paginate the results
        paginator = Paginator(manuscripts, 5)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'ccsrepo_app/adviser_manuscript.html', {
            'page_obj': page_obj,
            'manuscripts': manuscripts,
        })
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def adviser_review(request, manuscript_id):
    # Check if the user is an adviser
    if request.user.is_adviser:
        # Fetch the manuscript or return 404 if not found
        manuscript = get_object_or_404(Manuscript, id=manuscript_id)

        if request.method == "POST":
            # Get feedback and decision from POST data
            feedback = request.POST.get('feedback')
            decision = request.POST.get('decision')

            # Update manuscript feedback, status, and approval
            manuscript.feedback = feedback
            manuscript.status = "approved" if decision == "approve" else "rejected"
            manuscript.is_approved = True if decision == "approve" else False
            manuscript.publication_date = timezone.now() if decision == "approve" else None
            manuscript.save()

            # Redirect to the adviser manuscripts page
            return redirect('adviser_manuscript')

        # Render the review page with the manuscript data
        return render(request, 'ccsrepo_app/adviser_review.html', {'manuscript': manuscript})

    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

#----------------End Adviser System ------------------------/

#----------------Student System ------------------------/
@login_required(login_url='login')
def student_manuscripts_view(request):
    # Check if the user is authenticated and is a student
    if request.user.is_student:
        manuscripts = Manuscript.objects.filter(
            student=request.user, 
            upload_show=True
        ).order_by('-upload_date')

        paginator = Paginator(manuscripts, 10)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'ccsrepo_app/student_manuscript.html', {
            'page_obj': page_obj,
        })
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

def manuscript_detail_view(request, manuscript_id):
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)

    # Calculate the progress percentage if there are pages to process
    if manuscript.page_count > 0:
        progress_percentage = (manuscript.current_page_count / manuscript.page_count) * 100
    else:
        progress_percentage = 0

    # Add the calculated percentage to the context
    context = {
        'manuscript': manuscript,
        'progress_percentage': progress_percentage
    }

    return render(request, 'ccsrepo_app/manuscript_detail.html', context)

def extract_text_from_page(pdf_file, page_number):
    """
    Extract text from a specific page of a PDF file.
    Falls back to OCR if no text is extracted.
    """
    try:
        pdf_document = fitz.open(pdf_file.path)
        page = pdf_document.load_page(page_number - 1)  # Zero-indexed in PyMuPDF

        # Extract text using PyMuPDF (basic text extraction)
        ocr_text = page.get_text("text")

        # If no text is extracted, perform OCR
        if not ocr_text.strip():
            print(f"No text found on page {page_number}. Attempting OCR...")
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img)

        return ocr_text.strip()
    except Exception as e:
        print(f"Error processing page {page_number}: {e}")
        return ""

def continue_scanning(request, manuscript_id):
    """
    Processes the next 10 pages of the manuscript for OCR, 
    or fewer if fewer pages remain.
    """
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)

    # Determine the range of pages to process
    pages_to_process = min(10, manuscript.remaining_page)
    if pages_to_process <= 0:
        print("No pages left to process.")
        return redirect('manuscript_detail', manuscript_id=manuscript.id)

    for i in range(pages_to_process):
        page_number = manuscript.current_page_count + i + 1

        # Check if this page has already been processed
        if PageOCRData.objects.filter(manuscript=manuscript, page_num=page_number).exists():
            print(f"Skipping page {page_number}: already processed.")
            continue

        # Extract OCR text and save to the database
        ocr_text = extract_text_from_page(manuscript.pdf_file, page_number)
        if ocr_text:  # Only save if there is text
            try:
                PageOCRData.objects.create(
                    manuscript=manuscript,
                    page_num=page_number,
                    text=ocr_text
                )
                print(f"Processed and saved page {page_number}.")
            except IntegrityError:
                print(f"Duplicate entry detected for page {page_number}. Skipping.")
        else:
            print(f"No text extracted for page {page_number}. Skipping.")

    # Update manuscript's progress
    manuscript.current_page_count += pages_to_process
    manuscript.remaining_page = max(0, manuscript.page_count - manuscript.current_page_count)
    manuscript.save()

    print(f"Scanning completed for {pages_to_process} pages. Current page: {manuscript.current_page_count}")

    # Redirect back to the manuscript details page
    return redirect('manuscript_detail', manuscript_id=manuscript.id)

def faculty_continue_scanning(request, manuscript_id):
    """
    Processes the next 10 pages of the manuscript for OCR, 
    or fewer if fewer pages remain.
    """
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)

    # Determine the range of pages to process
    pages_to_process = min(1, manuscript.remaining_page)
    if pages_to_process <= 0:
        print("No pages left to process.")
        return redirect('faculty_detail', manuscript_id=manuscript.id)

    for i in range(pages_to_process):
        page_number = manuscript.current_page_count + i + 1

        # Check if this page has already been processed
        if PageOCRData.objects.filter(manuscript=manuscript, page_num=page_number).exists():
            print(f"Skipping page {page_number}: already processed.")
            continue

        # Extract OCR text and save to the database
        ocr_text = extract_text_from_page(manuscript.pdf_file, page_number)
        if ocr_text:  # Only save if there is text
            try:
                PageOCRData.objects.create(
                    manuscript=manuscript,
                    page_num=page_number,
                    text=ocr_text
                )
                print(f"Processed and saved page {page_number}.")
            except IntegrityError:
                print(f"Duplicate entry detected for page {page_number}. Skipping.")
        else:
            print(f"No text extracted for page {page_number}. Skipping.")

    # Update manuscript's progress
    manuscript.current_page_count += pages_to_process
    manuscript.remaining_page = max(0, manuscript.page_count - manuscript.current_page_count)
    manuscript.save()

    print(f"Scanning completed for {pages_to_process} pages. Current page: {manuscript.current_page_count}")

    # Redirect back to the manuscript details page
    return redirect('faculty_detail', manuscript_id=manuscript.id)

#----------------End Student System ------------------------/

#----------------Faculty System ------------------------/

@login_required(login_url='login')
def faculty_manuscripts_view(request):
    # Check if the user is an adviser or admin
    if request.user.is_adviser or request.user.is_admin:
        # Fetch manuscripts for the authenticated user
        manuscripts = Manuscript.objects.filter(
            adviser=request.user,  # Assuming adviser is related to Manuscript
            upload_show=True
        ).order_by('-publication_date')
        
        # Set up pagination
        paginator = Paginator(manuscripts, 5)  # Show 2 manuscripts per page
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        # Render the manuscripts page
        return render(request, 'ccsrepo_app/faculty_manuscript.html', {
            'page_obj': page_obj,
        })
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def faculty_detail_view(request, manuscript_id):
    # Check if the user is an adviser or an admin
    if request.user.is_adviser or request.user.is_admin:
        # Fetch the manuscript or return 404 if not found
        manuscript = get_object_or_404(Manuscript, id=manuscript_id)

        # Calculate the progress percentage if there are pages to process
        if manuscript.page_count > 0:
            progress_percentage = (manuscript.current_page_count / manuscript.page_count) * 100
        else:
            progress_percentage = 0

        # Add the calculated percentage to the context
        context = {
            'manuscript': manuscript,
            'progress_percentage': progress_percentage
        }

        # Render the faculty detail page
        return render(request, 'ccsrepo_app/faculty_detail.html', context)
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)
#----------------Faculty Upload System ------------------------/
@login_required(login_url='login')
def faculty_upload_manuscript(request):
    # Check if the user is an adviser or an admin
    if request.user.is_adviser or request.user.is_admin:
        if request.method == 'POST':
            # Get the uploaded file from the request
            pdf_file = request.FILES.get('pdf_file')

            if pdf_file:
                # Create a new Manuscript object
                manuscript = Manuscript(
                    pdf_file=pdf_file,
                    student=request.user,
                    abstracts="No abstract found"  # Default value if abstract is not extracted
                )
                manuscript.save()

                # Get the file path of the saved PDF
                pdf_file_path = manuscript.pdf_file.path

                # Process the manuscript if the file exists
                if os.path.exists(pdf_file_path):
                    try:
                        process_and_extract_manuscript_data(pdf_file_path, manuscript, max_abstract_pages=5, chunk_size=5)
                    except Exception as e:
                        # Log or print the error for debugging
                        print(f"Error processing PDF: {e}")

                # Redirect to the final confirmation page with the manuscript ID
                return redirect('faculty_final_page', manuscript_id=manuscript.id)

        # Render the upload page for GET requests
        return render(request, 'ccsrepo_app/faculty_upload_page.html')
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

@login_required(login_url='login')
def faculty_final_page(request, manuscript_id):
    if request.user.is_adviser or request.user.is_admin:
        manuscript = get_object_or_404(Manuscript, id=manuscript_id)
        errors = []

        if request.method == 'POST':
            # Get form data
            title = request.POST.get('title')
            abstracts = request.POST.get('abstracts')
            authors = request.POST.get('authors')
            year = request.POST.get('year')
            category_id = request.POST.get('category')
            manuscript_type_id = request.POST.get('manuscript_type')
            program_id = request.POST.get('program')

            # Validate year
            try:
                year = int(year)
                if year < 1900 or year > timezone.now().year:
                    errors.append(_("Please enter a valid year between 1900 and the current year."))
            except ValueError:
                errors.append(_("Year must be a valid number."))

            # Validate title uniqueness for `upload_show=True`
            if Manuscript.objects.filter(title=title, upload_show=True).exclude(id=manuscript.id).exists():
                errors.append(_("A manuscript with this title is already published. Please choose a different title."))

            # If there are validation errors, return them to the template
            if errors:
                categories = Category.objects.all()
                manuscript_types = ManuscriptType.objects.all()
                programs = Program.objects.all()
                return render(request, 'ccsrepo_app/faculty_final_page.html', {
                    'manuscript': manuscript,
                    'categories': categories,
                    'manuscript_types': manuscript_types,
                    'programs': programs,
                    'errors': errors,
                })

            # Update manuscript fields with the form data
            manuscript.title = title
            manuscript.abstracts = abstracts
            manuscript.authors = authors
            manuscript.year = year
            manuscript.category_id = category_id
            manuscript.manuscript_type_id = manuscript_type_id
            manuscript.program_id = program_id
            manuscript.adviser = request.user

            # Finalize and publish manuscript
            manuscript.publication_date = timezone.now()
            manuscript.status = 'approved'
            manuscript.upload_show = True

            manuscript.save()

            # **Keyword Extraction After Save**
            # Combine relevant fields for keyword extraction
            combined_text = f"{title} {abstracts} {authors}".lower()

            # Extract keywords based on tech-related terms
            extracted_keywords = extract_keywords_from_text(combined_text)

            # Save extracted keywords to the database
            existing_keywords = set(Keyword.objects.filter(manuscript=manuscript).values_list('keyword', flat=True))
            for keyword in extracted_keywords:
                if keyword not in existing_keywords:
                    Keyword.objects.create(manuscript=manuscript, keyword=keyword)

            # Redirect to the visitor search page
            return redirect('visitor_search_manuscripts')

        # Load choices for GET requests
        categories = Category.objects.all()
        manuscript_types = ManuscriptType.objects.all()
        programs = Program.objects.all()

        return render(request, 'ccsrepo_app/faculty_final_page.html', {
            'manuscript': manuscript,
            'categories': categories,
            'manuscript_types': manuscript_types,
            'programs': programs,
        })
    else:
        return render(request, 'unauthorized.html', status=403)

#----------------End Faculty Upload System ------------------------/

# ----------------Request Access System ------------------------/
def request_access(request, manuscript_id):
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)
    
    # Check if the user is the assigned student or already has access
    if request.user.is_student and request.user == manuscript.student:
        return redirect('view_pdf', manuscript_id=manuscript.id)

    # Check if an access request already exists for this student and manuscript
    existing_request = ManuscriptAccessRequest.objects.filter(
        manuscript=manuscript, student=request.user, status='pending'
    ).exists()

    if not existing_request:
        # Create a new access request if none exists
        ManuscriptAccessRequest.objects.create(
            manuscript=manuscript,
            student=request.user,
        )
    return redirect('visitor_manuscript_detail', manuscript_id=manuscript_id)


@login_required(login_url='login')
def manuscript_access_requests(request):
    # Check if the user is authenticated and is a student
    if request.user.is_adviser or request.user.is_admin or request.user.is_student:
        access_requests = ManuscriptAccessRequest.objects.filter(
        manuscript__adviser=request.user
    ).select_related('manuscript').order_by('-requested_at')

        paginator = Paginator(access_requests, 10)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'ccsrepo_app/manuscript_access_requests.html', {
            'page_obj': page_obj,
        })
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)

def manage_access_request(request):
    # Manage approval or denial of a request via a single endpoint
    if request.method == "POST":
        request_id = request.POST.get("request_id")
        action = request.POST.get("action")
        
        # Retrieve the access request and verify the adviser is responsible
        access_request = get_object_or_404(ManuscriptAccessRequest, id=request_id, manuscript__adviser=request.user)
        
        if action == "approve":
            # Approve and set the access duration
            access_request.approve(duration_days=7)
        elif action == "deny":
            # Deny the access request
            access_request.deny()
        
    return redirect("manuscript_access_requests")

@login_required(login_url='login')
def student_access_requests(request):
    if request.user.is_student:
        # Fetch access requests for the logged-in student
        access_requests = ManuscriptAccessRequest.objects.filter(student=request.user)

        # Set up pagination for the access requests
        paginator = Paginator(access_requests, 10)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'ccsrepo_app/student_access_requests.html', {
            'access_requests': page_obj.object_list,
            'page_obj': page_obj,
        })
    else:
        # Render the unauthorized page for unauthorized users
        return render(request, 'unauthorized.html', status=403)


def delete_unpublished_manuscripts(request):
    if request.method == 'POST':
        # Filter manuscripts where upload_show is False
        manuscripts_to_delete = Manuscript.objects.filter(upload_show=False)

        for manuscript in manuscripts_to_delete:
            # Delete related PageOCR records
            PageOCRData.objects.filter(manuscript=manuscript).delete()

            # Delete the PDF file from the filesystem
            if manuscript.pdf_file and os.path.isfile(manuscript.pdf_file.path):
                try:
                    os.remove(manuscript.pdf_file.path)
                    print(f"Deleted PDF file: {manuscript.pdf_file.path}")
                except Exception as e:
                    print(f"Failed to delete PDF file {manuscript.pdf_file.path}: {e}")

            # Delete the Manuscript itself
            manuscript.delete()
            print(f"Deleted manuscript: {manuscript.title}")

        messages.success(request, "Unpublished manuscripts have been successfully deleted.")
        return redirect('delete_unpublished_manuscripts')  # Redirect to your desired page after deletion
    return render(request, 'ccsrepo_app/delete_unpublished_manuscript.html')

# ----------------Indexing System ------------------------/
# Define the list of tech-related keywords
tech_related_keywords = [
    "machine learning", "artificial intelligence", "deep learning", "neural network", "data science",
    "python", "algorithm", "natural language processing", "computer vision", "big data", "cloud computing",
    "data analysis", "predictive modeling", "data mining", "automation", "robotics", "internet of things", 
    "AI", "IoT", "blockchain", "virtual reality", "augmented reality", "machine vision",
    "javascript", "node.js", "react", "angular", "vue.js", "express.js",
    "html", "css", "web development", "front-end", "back-end", "full-stack", "typescript", 
    "ruby on rails", "flutter", "swift", "kotlin", "c++", "c#", "rust", "bash", "php", "sql", "nosql",
    "docker", "kubernetes", "aws", "azure", "gcp", "git", "gitlab", "github", "jenkins", "ci/cd", "tensorflow", 
    "keras", "pytorch", "scikit-learn", "matplotlib", "pandas", "numpy", "opencv", "django", "flask", "spark"
]

def extract_keywords_from_text(text):
    """
    Extract exact matches for keywords from a predefined list, ensuring only full matches.
    """
    # Convert text to lowercase to ensure case-insensitivity
    text = text.lower()
    
    # Only add full matches of keywords, preventing partial matches like "Java" for "JavaScript"
    keywords_found = []
    for keyword in tech_related_keywords:
        # Ensure we match the exact word (not a part of other words)
        if re.search(r'\b' + re.escape(keyword) + r'\b', text):
            keywords_found.append(keyword)
    
    # Return the found keywords (no duplicates)
    return list(set(keywords_found))

def extract_keywords_after_keywords(text):
    """
    Extract keywords after 'KEYWORDS:' section in the text.
    This function is specifically designed to capture keywords that follow 'KEYWORDS:' directly.
    """
    keywords_section = re.search(r'keywords:\s*(.*)', text, re.IGNORECASE)
    if keywords_section:
        keywords_str = keywords_section.group(1)
        # Split the keywords by commas or newlines and clean them
        keywords = [keyword.strip().lower() for keyword in keywords_str.split(',')]
        return keywords
    return []

def clean_and_extract_after_keywords(text):
    """
    Clean and extract content after 'CHAPTER', 'EXECUTIVE SUMMARY', or 'KEYWORDS:' keywords.
    """
    patterns = [
        r'(?<=\bchapter\b)(.*)',  # Match after "CHAPTER"
        r'(?<=\bexecutive summary\b)(.*)',  # Match after "EXECUTIVE SUMMARY"
        r'(?<=\bkeywords:\b)(.*)',  # Match after "KEYWORDS:"
    ]
    
    # Check for matches and return the first match (prioritize CHAPTER, EXECUTIVE SUMMARY, or KEYWORDS: )
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()  # Return the content after the keyword
    
    return text  # If no match, return the original text

@login_required(login_url='login')
def view_pdf_manuscript(request, manuscript_id):
    if not (request.user.is_admin or request.user.is_adviser or request.user.is_student):
        return render(request, 'unauthorized.html', status=403)
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)
    
    # Safely assign a value to 'search_term' even if it's not in the GET request
    search_term = request.GET.get('search', '').strip()  # Default to an empty string if 'search' is not in the GET request

    # Prepare OCR data for the view
    if search_term:
        ocr_data = manuscript.ocr_data.filter(text__icontains=search_term).order_by('page_num')
        for page in ocr_data:
            highlighted_text = re.sub(
                f"({re.escape(search_term)})",
                r'<span class="highlight">\1</span>',
                page.text,
                flags=re.IGNORECASE
            )
            page.highlighted_text = mark_safe(highlighted_text)
    else:
        ocr_data = manuscript.ocr_data.all().order_by('page_num')
        for page in ocr_data:
            page.highlighted_text = page.text

    # Handle AJAX requests to return only OCR data
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        response_data = [{'page_num': page.page_num, 'highlighted_text': page.highlighted_text} for page in ocr_data]
        return JsonResponse({'ocr_data': response_data})

    # Collect matching page numbers for the search results
    matching_page_numbers = [page.page_num for page in ocr_data]

    return render(request, 'ccsrepo_app/view_pdf_manuscript.html', {
        'manuscript': manuscript,
        'ocr_data': ocr_data,
        'search_term': search_term,
        'matching_page_numbers': matching_page_numbers,
    })
# ----------------End Indexing System ------------------------/

# ----------------Visitor Site Flow System ------------------------/
def index_view(request):
    if request.method == 'GET' and 'q' in request.GET:
        search_query = request.GET.get('q')
        # Redirect to the search result page with the search query
        return redirect(f'/search/?q={search_query}')  # Use the exact URL of your search results page
    
    return render(request, 'index.html')

def visitor_search_manuscripts(request):
    # Get the search query if provided
    search_query = request.GET.get('q', '')

    # Split the search query into individual tags (words)
    tags = search_query.split()  # Split by whitespace to treat each term as a tag

    # Base queryset for manuscripts, filtering only approved manuscripts with non-null publication_date
    manuscripts = Manuscript.objects.filter(status='approved', publication_date__isnull=False)

    # Construct search filter using Q objects for multiple fields
    search_filter = Q()

    if search_query:
        # Loop through each tag and add it to the search filter
        for tag in tags:
            # Apply a filter to ensure the tag appears in all fields for a manuscript to be included
            search_filter &= (
                Q(title__icontains=tag) |
                Q(authors__icontains=tag) |
                Q(abstracts__icontains=tag) |
                Q(year__icontains=tag) |
                Q(adviser__first_name__icontains=tag) |
                Q(adviser__middle_name__icontains=tag) |
                Q(adviser__last_name__icontains=tag) |
                Q(program__name__icontains=tag) |
                Q(category__name__icontains=tag) |
                Q(manuscript_type__name__icontains=tag)
            )

    # Apply the search filter
    manuscripts = manuscripts.filter(search_filter)

    # Apply additional filters if selected
    program = None
    manuscript_type = None
    category = None
    year = None

    if request.GET.get('program'):
        program = Program.objects.get(id=request.GET['program'])
        manuscripts = manuscripts.filter(program_id=program.id)

    if request.GET.get('manuscript_type'):
        manuscript_type = ManuscriptType.objects.get(id=request.GET['manuscript_type'])
        manuscripts = manuscripts.filter(manuscript_type_id=manuscript_type.id)

    if request.GET.get('category'):
        category = Category.objects.get(id=request.GET['category'])
        manuscripts = manuscripts.filter(category_id=category.id)

    if request.GET.get('year'):  # Apply filter for the year
        year = request.GET.get('year')
        manuscripts = manuscripts.filter(publication_date__year=year)

    # Apply filters for status=approved and publication_date is not null
    manuscripts = manuscripts.filter(status='approved', publication_date__isnull=False)

    # Order manuscripts by publication_date in descending order
    manuscripts = manuscripts.order_by('-publication_date')

    # Annotate counts for programs, manuscript types, and categories for only approved manuscripts with a non-null publication_date
    programs = Program.objects.annotate(
        manuscript_count=Count(
            'manuscript', 
            filter=Q(manuscript__status='approved') & Q(manuscript__publication_date__isnull=False)
        )
    )
    manuscript_types = ManuscriptType.objects.annotate(
        manuscript_count=Count(
            'manuscript', 
            filter=Q(manuscript__status='approved') & Q(manuscript__publication_date__isnull=False)
        )
    )
    categories = Category.objects.annotate(
        manuscript_count=Count(
            'manuscript', 
            filter=Q(manuscript__status='approved') & Q(manuscript__publication_date__isnull=False)
        )
    )

    # Group manuscripts by year field and count approved manuscripts per year, ensuring publication_date is not null
    manuscript_years = (
        manuscripts.values('year')
        .annotate(count=Count('id', filter=Q(status='approved') & Q(publication_date__isnull=False)))
        .order_by('-year')  # Sort by year descending
    )

    # Pagination setup
    paginator = Paginator(manuscripts, 10)  # Show 10 manuscripts per page
    page_number = request.GET.get('page')
    manuscripts_page = paginator.get_page(page_number)

    # Pass context to the template, including selected filter names
    context = {
        'manuscripts': manuscripts_page,
        'programs': programs,
        'manuscript_types': manuscript_types,
        'categories': categories,
        'search_query': search_query,
        'selected_program': program,
        'selected_manuscript_type': manuscript_type,
        'selected_category': category,
        'selected_year': year,  # Include the selected year filter
        'manuscript_years': manuscript_years,  # Include year aggregation
    }
    return render(request, 'visitor_search_result.html', context)

def visitor_manuscript_detail(request, manuscript_id):
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)
    
    # Replace commas with <br> for the authors
    authors_with_br = manuscript.authors.replace(',', '<br>')

    # Check if the user is authenticated and determine roles
    is_authenticated = request.user.is_authenticated
    is_student = is_authenticated and request.user == manuscript.student
    is_admin = is_authenticated and getattr(request.user, 'is_admin', False)
    is_adviser = is_authenticated and getattr(request.user, 'is_adviser', False)

    # Check if the user has an approved access request
    access_request = None
    has_pending_request = False

    if is_authenticated:
        access_request = ManuscriptAccessRequest.objects.filter(
            manuscript=manuscript,
            student=request.user,
            status='approved',
            access_start_date__lte=timezone.now(),
            access_end_date__gte=timezone.now()
        ).first()
        
        # Check if the user has a pending access request
        has_pending_request = ManuscriptAccessRequest.objects.filter(
            manuscript=manuscript,
            student=request.user,
            status='pending'
        ).exists()
    # Set has_access to True if the user is the student, adviser, admin, or has an approved request
    has_access = is_student or is_adviser or is_admin or (access_request is not None)
    ip_address = request.META.get('HTTP_X_FORWARDED_FOR', None)
    if ip_address:
        ip_address = ip_address.split(',')[0]  # If there are multiple IPs, take the first one
    else:
        ip_address = request.META.get('REMOTE_ADDR')  # Fallback to the direct REMOTE_ADDR

    # Check if the IP address has already viewed the manuscript
    if not ManuscriptView.objects.filter(manuscript=manuscript, ip_address=ip_address).exists():
        # Increment views
        manuscript.views += 1
        manuscript.save()

        # Create a new ManuscriptView record to track the view
        ManuscriptView.objects.create(manuscript=manuscript, ip_address=ip_address)
    return render(request, 'visitor_manuscript_detail.html', {
        'manuscript': manuscript,
        'authors_with_br': authors_with_br,
        'has_access': has_access,
        'has_pending_request': has_pending_request,
    })

def delete_manuscript(request, manuscript_id):
    manuscript = get_object_or_404(Manuscript, id=manuscript_id)
    
    # Check if the user has permission to delete the manuscript
    if request.user.is_admin:
        manuscript.delete()
    else:
        return render(request, 'unauthorized.html', status=403)
    
    return redirect('visitor_search_manuscripts')

from docx import Document
from django.http import HttpResponse
from django.db.models import Sum
@login_required(login_url='login')
def generate_reports(request):
    if request.method == 'POST':
        # Get the selected filters
        adviser_id = request.POST.get('adviser')
        program_id = request.POST.get('program')
        category_id = request.POST.get('category')
        type_id = request.POST.get('type')  # Manuscript type
        year = request.POST.get('year')
        publication_year = request.POST.get('publication_year')

        # Query the database with the selected filters
        manuscripts = Manuscript.objects.filter(status='approved', publication_date__isnull=False)

        # Apply filters based on the selected options
        if adviser_id:
            manuscripts = manuscripts.filter(adviser_id=adviser_id)
        if program_id:
            manuscripts = manuscripts.filter(program_id=program_id)
        if category_id:
            manuscripts = manuscripts.filter(category_id=category_id)
        if type_id:
            try:
                type_id = int(type_id)
                manuscripts = manuscripts.filter(manuscript_type_id=type_id)
            except ValueError:
                print(f"Invalid manuscript type ID: {type_id}")
        if year:
            manuscripts = manuscripts.filter(year=year)
        if publication_year:
            manuscripts = manuscripts.filter(publication_date__year=publication_year)

        # For preview: create HTML table summary with your custom CSS class
        preview_html = """
        <table class="manuscript-table">
            <thead>
                <tr>
                    <th>Title</th>
                    <th>Adviser</th>
                    <th>Category</th>
                    <th>Program</th>
                    <th>Type</th>
                    <th>Year</th>
                    <th>Authors</th>
                </tr>
            </thead>
            <tbody>
        """
        
        # Add rows for each manuscript
        for manuscript in manuscripts:
            authors = manuscript.authors if manuscript.authors else 'N/A'
            preview_html += f"""
                <tr>
                    <td>{manuscript.title}</td>
                    <td>{manuscript.adviser.first_name} {manuscript.adviser.last_name}</td>
                    <td>{manuscript.category.name if manuscript.category else 'N/A'}</td>
                    <td>{manuscript.program.name if manuscript.program else 'N/A'}</td>
                    <td>{manuscript.manuscript_type.name if manuscript.manuscript_type else 'N/A'}</td>
                    <td>{manuscript.year}</td>
                    <td>{authors}</td>
                </tr>
            """
        
        preview_html += "</tbody></table>"

        # Return the preview HTML to be displayed on the frontend
        return JsonResponse({
            'preview_html': preview_html
        })

    # Render the filter form (unchanged)
    advisers = CustomUser.objects.filter(is_adviser=True)
    programs = Program.objects.all()
    categories = Category.objects.all()
    types = ManuscriptType.objects.all()
    years = Manuscript.objects.filter(year__isnull=False).values('year').distinct()
    publication_years = Manuscript.objects.filter(publication_date__isnull=False).values('publication_date__year').distinct()

    return render(request, 'ccsrepo_app/generate_reports.html', {
        'advisers': advisers,
        'programs': programs,
        'categories': categories,
        'types': types,
        'years': years,
        'publication_years': publication_years,
    })

from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
@login_required(login_url='login')
def download_report(request):
    # Get filter parameters from the query string
    adviser_id = request.GET.get('adviser')
    program_id = request.GET.get('program')
    category_id = request.GET.get('category')
    type_id = request.GET.get('type')
    year = request.GET.get('year')
    publication_year = request.GET.get('publication_year')

    # Start the query to get approved manuscripts with valid publication dates
    manuscripts = Manuscript.objects.filter(status='approved', publication_date__isnull=False)

    # Apply filters based on the selected options
    if adviser_id:
        manuscripts = manuscripts.filter(adviser_id=adviser_id)
    if program_id:
        manuscripts = manuscripts.filter(program_id=program_id)
    if category_id:
        manuscripts = manuscripts.filter(category_id=category_id)
    if type_id:  # Ensure the type filter is applied
        manuscripts = manuscripts.filter(manuscript_type_id=type_id)
    if year:
        manuscripts = manuscripts.filter(year=year)
    if publication_year:
        manuscripts = manuscripts.filter(publication_date__year=publication_year)

    # Debugging (Optional): Check filtered manuscripts
    print("Manuscripts after filtering:", manuscripts.count())

    # Generate the report as a Word document
    doc = Document()
    doc.add_heading('Generated Report', 0)
    doc.add_paragraph(f"Report generated on: {timezone.now()}")
    doc.add_paragraph("\n")

    # Add a summary
    total_views = manuscripts.aggregate(Sum('views'))['views__sum'] or 0
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f"Total manuscripts generated: {manuscripts.count()}")
    doc.add_paragraph(f"Total views across all manuscripts: {total_views}")

    # Create a table in the Word document
    doc.add_heading('Manuscripts', level=1)

    # Define the headers for the table
    headers = ['Title', 'Adviser', 'Category', 'Program', 'Type', 'Year', 'Authors']
    table = doc.add_table(rows=1, cols=len(headers))  # Create a table with the correct number of columns

    # Add header cells with a normal style (bold text, no color)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align header text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True  # Make header text bold

    # Add rows for each manuscript and handle alternating row colors
    for index, manuscript in enumerate(manuscripts):
        row_cells = table.add_row().cells
        row_cells[0].text = manuscript.title
        row_cells[1].text = f"{manuscript.adviser.first_name} {manuscript.adviser.last_name}"
        row_cells[2].text = manuscript.category.name if manuscript.category else 'N/A'
        row_cells[3].text = manuscript.program.name if manuscript.program else 'N/A'
        row_cells[4].text = manuscript.manuscript_type.name if manuscript.manuscript_type else 'N/A'
        row_cells[5].text = str(manuscript.year)
        row_cells[6].text = manuscript.authors if manuscript.authors else 'N/A'

        # Style for alternating row colors
        if index % 2 == 0:  # Even index - light yellow
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
                cell._element.xpath('.//w:shd')[0].set(qn('w:fill'), "FAF1E6")  # Light background for even rows
        else:  # Odd index - white
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
                cell._element.xpath('.//w:shd')[0].set(qn('w:fill'), "FFFFFF")  # White background for odd rows

    # Save the document to the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="report.docx"'
    doc.save(response)

    return response

def get_filtered_options(request, adviser_id):
    manuscripts = Manuscript.objects.filter(
        status='approved',
        publication_date__isnull=False
    )
    if adviser_id != "all":
        manuscripts = manuscripts.filter(adviser_id=adviser_id)

    programs = Program.objects.filter(manuscript__in=manuscripts).distinct().values('id', 'name')
    categories = Category.objects.filter(manuscript__in=manuscripts).distinct().values('id', 'name')
    types = ManuscriptType.objects.filter(manuscript__in=manuscripts).distinct().values('id', 'name')
    years = manuscripts.values('year').distinct()
    publication_years = manuscripts.values('publication_date__year').distinct()

    data = {
        'programs': list(programs),
        'categories': list(categories),
        'types': list(types),
        'years': list(years),
        'publication_years': list(publication_years),
    }
    return JsonResponse(data)